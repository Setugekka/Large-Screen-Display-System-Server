# encoding=utf8
import sys
reload(sys)
sys.setdefaultencoding('utf8')
import os
from flask import Blueprint,jsonify,abort,make_response,request,send_file,send_from_directory
from docx import Document
from docx.shared import Pt,Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
report_blueprint = Blueprint(
    'report',
    __name__,
    url_prefix="/report"
)
class Report:
    title=''
    author=''
    content=''
    department=''
    contact=''
class GenerateReport:
    newReport = Report()
    def __init__(self,__report):
        self.newReport=__report
    def generateRep(self):
        document = Document()
        document.add_heading(self.newReport.title,level=0)
        author = document.add_paragraph(self.newReport.author)
        author_format = author.paragraph_format
        author_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        contact = document.add_paragraph(self.newReport.contact)
        contact_format = contact.paragraph_format
        contact_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        department = document.add_paragraph(self.newReport.department)
        department_format = department.paragraph_format
        department_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph(self.newReport.content)
        document.save(self.newReport.title+'.docx')
tasks=[
    {'id':1,
    'title':u'买菜',
    'description':u'牛奶,奶酪,披萨,水果',
    'done':False},
    {'id':3,
    'title':u'学习python',
    'description':u'需要在网上找个合适的python教程',
    'done':False}

]
@report_blueprint.route('/todo/api/v1.0/tasks',methods=['GET'])
def get_tasks():
    return jsonify(tasks)
@report_blueprint.route('/todo/api/v1.0/tasks/<int:task_id>',methods=['GET'])
def get_task(task_id):
    task = filter(lambda t:t['id']==task_id,tasks)
    if len(task) == 0:
        abort(404)
    return jsonify({'data':task[0]})

@report_blueprint.errorhandler(404)
def not_found(error):
    return make_response(jsonify({'error':'Not Found'}),404)

@report_blueprint.route('/todo/api/v1.0/tasks',methods=['POST'])
def create_task():
    if not request.json or not 'title' in request.json:
        abort(400)
    task = {
        "id":tasks[-1]['id']+1,
        'title':request.json['title'],
        'description':request.json.get("description",''),
        'done':False
    }
    tasks.append(task)
    return jsonify(task)
@report_blueprint.route('/testapi',methods=['POST'])
def test():
    newRep = Report()
    # print request.get_json()
    # print request.json['name']
    # print request.json['title']
    # print request.json['content']
    newRep.author = request.json['name']
    newRep.title=request.json['title']
    newRep.content=request.json['content']
    genRep=GenerateReport(newRep)
    genRep.generateRep()
    return jsonify({'success':'true'})
@report_blueprint.route('/GenerateReportApi',methods=['POST']) #报表自动生成api
def generateReport():
    newRep = Report()
    # print request.get_json()
    # print request.json['name']
    # print request.json['title']
    # print request.json['content']
    newRep.author = request.json['name']
    newRep.title=request.json['title']
    newRep.content=request.json['content']
    newRep.department=request.json['department']
    newRep.contact=request.json['contact']
    genRep=GenerateReport(newRep)
    genRep.generateRep()
    return jsonify({'success':'true'})

#文件下载api
@report_blueprint.route('/download/<filename>',methods=['GET'])
def download_file(filename):
    file_name = filename.decode("latin-1") #将utf-8解码为latin-1
    directory = os.getcwd()
    response = make_response(send_from_directory(directory,filename,as_attachment=True))
    response.headers["Content-Disposition"] = "attachment; filename="+file_name
    return response



